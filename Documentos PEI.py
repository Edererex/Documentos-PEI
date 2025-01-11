import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import pandas as pd
import os
import sys
from docx import Document
import threading
import json
from datetime import datetime

# Caminho para o arquivo de licença
LICENCA_FILE_PATH = 'licenca.json'

# Função para carregar a licença armazenada localmente (se existir)
def carregar_licenca_local():
    if os.path.exists(LICENCA_FILE_PATH):
        try:
            with open(LICENCA_FILE_PATH, 'r') as f:
                return json.load(f)
        except json.JSONDecodeError:
            return None
    return None

# Função para salvar a licença localmente
def salvar_licenca_local(licenca):
    with open(LICENCA_FILE_PATH, 'w') as f:
        json.dump(licenca, f)

# Função para verificar se a licença é válida
def validar_licenca(codigo_entrada, licenca):
    codigo_valido = codigo_entrada == licenca["codigo"]
    ainda_valido = datetime.now() <= datetime.strptime(licenca["data_expiracao"], '%Y-%m-%d %H:%M:%S.%f')
    
    if codigo_valido and ainda_valido:
        print("Licença válida! Bem-vindo!")
        return True
    elif not codigo_valido:
        print("Código de licença inválido.")
        return False
    else:
        print("Licença expirada.")
        return False

class Application(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Documentos PEI")
        self.geometry("500x450")
        self.excel_file_path = None
        self.word_template_path = None
        self.excel_data = None
        self.additional_params = []  # Lista para armazenar os parâmetros adicionais

        # Solicita o código de licença se necessário
        if not self.validar_licenca():
            self.destroy()  # Fecha a aplicação se a licença for inválida
            return
        
        self.create_widgets()

    # Função para validar a licença e solicitá-la apenas se necessário
    def validar_licenca(self):
        # Carrega a licença local (se existir)
        licenca = carregar_licenca_local()

        # Se a licença existe e é válida, não pede o código
        if licenca:
            if validar_licenca(licenca["codigo"], licenca):
                return True

        # Se a licença não for válida ou não existir, solicita uma nova
        return self.solicitar_nova_licenca()

    # Função para solicitar uma nova licença do usuário
    def solicitar_nova_licenca(self):
        # Solicita o código de licença
        codigo_usuario = self.solicitar_codigo_licenca()

        # Aqui você pode simular uma verificação ou validação desse código com o servidor ou gerar uma nova licença localmente
        nova_licenca = {
            "codigo": codigo_usuario,
            "data_expiracao": (datetime.now() + pd.DateOffset(days=30)).strftime('%Y-%m-%d %H:%M:%S.%f')  # Licença válida por 30 dias
        }

        # Salva a nova licença
        salvar_licenca_local(nova_licenca)

        # Valida a nova licença gerada
        return validar_licenca(codigo_usuario, nova_licenca)

    # Função para abrir uma janela solicitando o código de licença
    def solicitar_codigo_licenca(self):
        janela_licenca = tk.Toplevel(self)
        janela_licenca.title("Verificação de Licença")
        janela_licenca.geometry("300x150")
        
        tk.Label(janela_licenca, text="Insira seu código de licença:").pack(pady=10)
        
        codigo_entry = tk.Entry(janela_licenca)
        codigo_entry.pack(pady=10)

        def confirmar_licenca():
            self.codigo_usuario = codigo_entry.get()  # Salva o código na instância da classe
            janela_licenca.destroy()  # Fecha a janela após a inserção do código
        
        botao_confirmar = tk.Button(janela_licenca, text="Confirmar", command=confirmar_licenca)
        botao_confirmar.pack(pady=10)
        
        self.wait_window(janela_licenca)  # Espera até que a janela seja fechada
        return self.codigo_usuario  # Retorna o código de licença fornecido

    def create_widgets(self):
        self.create_labels_and_entries()
        self.create_buttons()

    def create_labels_and_entries(self):
        labels = ["Disciplina", "Bimestre", "Ano/Série", "Ciclo", "Aula", "Quantidade de Aulas"]
        self.entries = {}

        for i, label_text in enumerate(labels):
            tk.Label(self, text=f"{label_text}:").grid(row=i, column=0, padx=10, pady=10)
            self.entries[label_text] = ttk.Combobox(self)
            self.entries[label_text].grid(row=i, column=1, padx=10, pady=10)

        # Definir valores padrão para os campos
        self.entries["Bimestre"].config(values=["1°", "2°", "3°", "4°"])
        self.entries["Ano/Série"].config(values=["1° ano", "2° ano", "3° ano", "4° ano", "5° ano", "6° ano", "7° ano", "8° ano", "9° ano", "1ª série", "2ª série", "3ª série"])
        self.entries["Ciclo"].config(values=["Anos Iniciais", "Anos Finais", "Ensino Médio"])
        self.entries["Aula"].config(values=list(range(1, 60)))  # Assume até 50 aulas
        self.entries["Quantidade de Aulas"].config(values=list(range(1, 60)))  # Assume até 10 aulas por vez

    def create_buttons(self):
        self.excel_button = tk.Button(self, text="Selecionar Arquivo Excel", command=self.select_excel_file)
        self.excel_button.grid(row=6, column=0, padx=10, pady=10)

        self.excel_label = tk.Label(self, text="", fg="blue")
        self.excel_label.grid(row=6, column=1, padx=10, pady=10)

        self.word_button = tk.Button(self, text="Selecionar Modelo Word", command=self.select_word_template)
        self.word_button.grid(row=7, column=0, padx=10, pady=10)

        self.word_label = tk.Label(self, text="", fg="blue")
        self.word_label.grid(row=7, column=1, padx=10, pady=10)

        self.generate_button = tk.Button(self, text="Gerar Documento", command=self.generate_document, bg="green", fg="white")
        self.generate_button.grid(row=8, column=0, columnspan=2, padx=10, pady=20)
    def select_excel_file(self):
        self.excel_file_path = filedialog.askopenfilename(title="Selecione o arquivo Excel", filetypes=[("Excel files", "*.xlsx")])
        if self.excel_file_path:
            self.excel_label.config(text=self.excel_file_path.split("/")[-1])
            self.load_excel_data_thread()

    def load_excel_data_thread(self):
        self.excel_button.config(state=tk.DISABLED)
        self.excel_label.config(text="Carregando...")
        threading.Thread(target=self.load_excel_data, daemon=True).start()

    def load_excel_data(self):
        try:
            self.excel_data = pd.read_excel(self.excel_file_path, sheet_name=None, header=1)
            disciplinas = list(self.excel_data.keys())
            self.after(0, self.update_ui_after_excel_load, disciplinas)
        except FileNotFoundError:
            self.after(0, lambda: messagebox.showerror("Erro", "Arquivo Excel não encontrado."))
        except pd.errors.EmptyDataError:
            self.after(0, lambda: messagebox.showerror("Erro", "O arquivo Excel está vazio."))
        except Exception as e:
            self.after(0, lambda: messagebox.showerror("Erro", f"Erro ao carregar o arquivo Excel: {str(e)}"))
        finally:
            self.after(0, lambda: self.excel_button.config(state=tk.NORMAL))

    def update_ui_after_excel_load(self, disciplinas):
        self.entries["Disciplina"].config(values=disciplinas)
        self.entries["Disciplina"].current(0)
        self.excel_label.config(text=self.excel_file_path.split("/")[-1])

    def select_word_template(self):
        self.word_template_path = filedialog.askopenfilename(title="Selecione o modelo Word", filetypes=[("Word files", "*.docx")])
        if self.word_template_path:
            self.word_label.config(text=self.word_template_path.split("/")[-1])
            
            # Verifica o número de tabelas no template Word
            doc = Document(self.word_template_path)
            num_tables = len(doc.tables)
            
            if num_tables > 1:
                response = messagebox.askyesno("Várias Tabelas Detectadas", 
                                               f"O documento contém {num_tables} tabelas. Deseja preencher todas as tabelas?")
                if response:
                    self.collect_additional_parameters(num_tables - 1)

    def collect_additional_parameters(self, num_extra_tables):
        self.additional_params = []

        # Janela para parâmetros adicionais
        additional_window = tk.Toplevel(self)
        additional_window.title("Parâmetros Adicionais")
    
        tk.Label(additional_window, text="Preencha os parâmetros para as tabelas adicionais").pack(pady=10)
    
        for i in range(num_extra_tables):
            frame = tk.Frame(additional_window)
            frame.pack(pady=10)
        
            params = {}
            tk.Label(frame, text=f"Tabela {i+2}:").grid(row=0, column=0, columnspan=2, pady=5)
        
            for j, label_text in enumerate(self.entries.keys()):
                tk.Label(frame, text=f"{label_text}:").grid(row=j+1, column=0, padx=10, pady=5)
                entry = ttk.Combobox(frame)
                entry.grid(row=j+1, column=1, padx=10, pady=5)
                entry.config(values=self.entries[label_text].cget('values'))  # Copia as mesmas opções do combobox principal
                params[label_text] = entry
        
            self.additional_params.append(params)

        def save_additional_params():
            # Armazena os valores dos comboboxes antes de fechar a janela
            for params in self.additional_params:
                for key, entry in params.items():
                    params[key] = entry.get()
        
            additional_window.destroy()  # Fecha a janela somente depois de armazenar os valores

        # Botão para confirmar e salvar os parâmetros
        tk.Button(additional_window, text="Confirmar", command=save_additional_params).pack(pady=10)

    def generate_document(self):
        # Coleta os valores dos parâmetros principais
        params = {key: entry.get() for key, entry in self.entries.items()}

        # Valida os parâmetros principais
        if any(value == "" for value in params.values()):
            messagebox.showerror("Erro", "Todos os campos devem ser preenchidos.")
            return

        try:
            params['Aula'] = int(params['Aula'])
            params['Quantidade de Aulas'] = int(params['Quantidade de Aulas'])
        except ValueError:
            messagebox.showerror("Erro", "Aula e Quantidade de Aulas devem ser números inteiros.")
            return

        if not self.excel_file_path or not self.word_template_path:
            messagebox.showwarning("Aviso", "Por favor, selecione os arquivos necessários.")
            return

        # Inicia a geração do documento em uma nova thread
        self.process_data_and_generate_document_thread(**params)

    def process_data_and_generate_document_thread(self, **params):
        self.generate_button.config(state=tk.DISABLED)
        progress_window = self.create_progress_window()
        threading.Thread(target=self.process_data_and_generate_document, args=(progress_window, params), daemon=True).start()

    def create_progress_window(self):
        progress_window = tk.Toplevel(self)
        progress_window.title("Gerando Documento")
        progress_window.geometry("300x100")
        
        tk.Label(progress_window, text="Gerando documento...").pack(pady=10)
        progress_bar = ttk.Progressbar(progress_window, mode='indeterminate')
        progress_bar.pack(pady=10)
        progress_bar.start()
        
        return progress_window

    def process_data_and_generate_document(self, progress_window, params):
        try:
            doc = Document(self.word_template_path)
            num_tables = len(doc.tables)
        
            # Preenche a primeira tabela com os parâmetros principais
            self.fill_table(doc.tables[0], params)
        
            # Preenche as tabelas adicionais se existirem
            for i, additional_params in enumerate(self.additional_params):
                if i + 1 < num_tables:  # Garante que não exceda o número de tabelas
                    self.fill_table(doc.tables[i + 1], additional_params)
        
            # Salvamento do documento
            file_options = {
                'defaultextension': '.docx',
                'filetypes': [('Word Document', '*.docx')],
                'initialfile': f"Plano_Aula_{params['Disciplina']}_{params['Bimestre']}_{params['Ano/Série']}_{params['Aula']}.docx",
                'title': "Salvar Documento"
            }
            save_path = filedialog.asksaveasfilename(**file_options)

            if save_path:
                doc.save(save_path)
                self.after(0, progress_window.destroy)
                self.after(0, lambda: messagebox.showinfo("Sucesso", f"Documento gerado e salvo como {save_path}"))
            else:
                self.after(0, progress_window.destroy)
                self.after(0, lambda: messagebox.showinfo("Cancelado", "O salvamento foi cancelado."))

        except Exception as e:
            error_message = str(e)  # Captura o erro em uma variável
            self.after(0, progress_window.destroy)
            self.after(0, lambda: messagebox.showerror("Erro", f"Ocorreu um erro ao gerar o documento: {error_message}"))
        finally:
            self.after(0, lambda: self.generate_button.config(state=tk.NORMAL))

    def fill_table(self, table, params):
        # Função para preencher uma tabela específica com base nos parâmetros fornecidos
        word_headers = [cell.text.strip() for cell in table.rows[0].cells]
        column_mapping = {}
        
        for col in self.excel_data[params["Disciplina"]].columns:
            for word_header in word_headers:
                if col.lower() == word_header.lower():
                    column_mapping[word_header] = col
                    break

        filtered_data = self.filter_data(params)
        
        for i, (idx, row_data) in enumerate(filtered_data.iterrows()):
            if i >= len(table.rows) - 1:
                table.add_row()

            row = table.rows[i + 1]
            for cell_idx, cell in enumerate(row.cells):
                header_text = word_headers[cell_idx]
                if header_text in column_mapping:
                    cell_value = row_data[column_mapping[header_text]]
                    cell.text = str(cell_value)

    def filter_data(self, params):
        # Filtra os dados do Excel com base nos parâmetros fornecidos
        sheet_data = self.excel_data[params["Disciplina"]]

        # Converte as colunas para os tipos corretos
        sheet_data['Ciclo'] = sheet_data['Ciclo'].astype(str)
        sheet_data['Ano/Série'] = sheet_data['Ano/Série'].astype(str)
        sheet_data['Bimestre'] = sheet_data['Bimestre'].astype(str)

        # Converte 'Aula' para numérico e remove valores não numéricos
        sheet_data['Aula'] = pd.to_numeric(sheet_data['Aula'], errors='coerce')
        sheet_data = sheet_data.dropna(subset=['Aula'])
        sheet_data['Aula'] = sheet_data['Aula'].astype(int)

        # Certifique-se de que os parâmetros também são do tipo correto
        params['Aula'] = int(params['Aula'])
        params['Quantidade de Aulas'] = int(params['Quantidade de Aulas'])
        params['Ciclo'] = str(params['Ciclo'])
        params['Ano/Série'] = str(params['Ano/Série'])
        params['Bimestre'] = str(params['Bimestre'])

        filtered_data = sheet_data[
            (sheet_data['Ciclo'] == params['Ciclo']) &
            (sheet_data['Ano/Série'] == params['Ano/Série']) &
            (sheet_data['Bimestre'] == params['Bimestre']) &
            (sheet_data['Aula'] >= params['Aula']) &
            (sheet_data['Aula'] < params['Aula'] + params['Quantidade de Aulas'])
        ]
        
        return filtered_data

if __name__ == "__main__":
    app = Application()
    app.mainloop()